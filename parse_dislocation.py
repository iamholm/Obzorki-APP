"""
parse_dislocation.py — Парсер файла дислокации (участковые уполномоченные).

Структура таблицы (5 колонок):
  C0: № о/п
  C1: № и адрес УПП
  C2: № адм. участка
  C3: Звание, должность, ФИО сотрудника
  C4: Обслуживаемые объекты (адреса)

Строки 0 и 1 — заголовки, данные начинаются со строки 2.
"""

import re
from docx import Document

import db

# Паттерн для поиска «УУП» или «старший УУП» в строке сотрудника
_POS_PAT   = re.compile(r'(?:старший\s+)?УУП', re.IGNORECASE)
_RANK_WORD = re.compile(
    r'\b(?:капитан|майор|лейтенант|полковник|подполковник|'
    r'полиции|ст\.|старший)\b', re.IGNORECASE)


def _parse_officer_cell(text: str) -> dict:
    """
    Разбирает ячейку «Звание, должность, ФИО» на поля.
    Возможные форматы:
      «Ст. лейтенант полиции\\nУУП\\nФамилия Имя Отчество»
      «ВАКАНСИЯ\\nмайор полиции\\nстарший УУП\\nФамилия Имя Отчество»
      «майор полиции старший УУП\\nФамилия Имя Отчество»
    """
    lines = [l.strip() for l in text.split('\n') if l.strip()]

    is_vacancy = bool(lines and lines[0].upper() == 'ВАКАНСИЯ')
    if is_vacancy:
        lines = lines[1:]

    if not lines:
        return {'rank': '', 'position': '', 'fio': '', 'is_vacancy': is_vacancy}

    # Ищем строку с «УУП» (позиция)
    pos_idx = next(
        (i for i, l in enumerate(lines) if _POS_PAT.search(l)),
        None)

    if pos_idx is None:
        # УУП не найден — пытаемся разделить по наличию слов звания:
        # строки со словами звания → rank, остальные → ФИО
        rank_lines, fio_lines = [], []
        fio_started = False
        for line in lines:
            if not fio_started and _RANK_WORD.search(line):
                rank_lines.append(line)
            else:
                fio_started = True
                fio_lines.append(line)
        return {
            'rank':     ' '.join(rank_lines).strip(),
            'position': 'УУП' if rank_lines else '',   # позиция по умолчанию
            'fio':      re.sub(r'\s+', ' ', ' '.join(fio_lines)).strip(),
            'is_vacancy': is_vacancy,
        }

    pos_line = lines[pos_idx]
    m = _POS_PAT.search(pos_line)

    # Часть строки до «УУП» — продолжение звания (напр. «майор полиции»)
    rank_tail = pos_line[:m.start()].strip()
    # Позиция = «УУП» или «старший УУП»
    position = pos_line[m.start():].strip()

    # Звание = строки ДО строки с УУП + хвост из той же строки
    rank_parts = lines[:pos_idx]
    if rank_tail:
        rank_parts = rank_parts + [rank_tail]
    rank = ' '.join(rank_parts).strip()

    # ФИО = строки ПОСЛЕ строки с УУП
    fio = re.sub(r'\s+', ' ', ' '.join(lines[pos_idx + 1:])).strip()

    return {'rank': rank, 'position': position, 'fio': fio, 'is_vacancy': is_vacancy}


def parse_and_save(doc_path: str) -> list:
    """
    Парсит файл дислокации, сохраняет сотрудников в БД.
    Возвращает список dict-ов с данными сотрудников.
    """
    doc = Document(doc_path)
    if not doc.tables:
        raise ValueError('В файле нет таблиц')

    table = doc.tables[0]
    officers = []

    for r_idx, row in enumerate(table.rows):
        if r_idx < 2:           # строки 0 и 1 — заголовки
            continue

        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 4:
            continue

        raw_upp      = cells[1] if len(cells) > 1 else ''
        raw_district = cells[2] if len(cells) > 2 else ''
        raw_officer  = cells[3]
        raw_addresses = cells[4] if len(cells) > 4 else ''

        if not raw_officer:
            continue

        parsed = _parse_officer_cell(raw_officer)
        if not parsed['fio']:
            continue

        # Извлекаем «УПП-NN»
        m = re.search(r'УПП-\d+', raw_upp, re.IGNORECASE)
        upp = m.group(0) if m else raw_upp.split('\n')[0].strip()

        # Номер участка: ячейка может содержать артефакт объединённой строки выше
        # (напр. «3\n48» или «3   48» — «3» это отдел, «48» — фактический участок).
        # Разбиваем на все токены и берём последний.
        _d_tokens = re.split(r'[\s]+', raw_district.strip())
        _d_tokens = [t for t in _d_tokens if t]
        district = _d_tokens[-1] if _d_tokens else ''

        officers.append({
            'fio':        parsed['fio'],
            'rank':       parsed['rank'],
            'position':   parsed['position'],
            'upp':        upp,
            'district':   district,
            'addresses':  re.sub(r'\s+', ' ', raw_addresses),
            'is_vacancy': int(parsed['is_vacancy']),
            'source_file': doc_path,
        })

    db.init_db()
    db.save_officers(officers)
    return officers
