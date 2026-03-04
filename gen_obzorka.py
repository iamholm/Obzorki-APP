"""
gen_obzorka.py — Библиотека генерации обзорных справок УИИ.

Экспортирует: generate_one(), _match_officer(), _officer_label(),
              CHAR_TEXTS, CONN_TEXTS, QUARTER_ACC, CHAR_OPTIONS,
              грамматические утилиты (_rank_to_instr, _fio_display, …)
"""

import os
import re
import copy
import random

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Шаблон ─────────────────────────────────────────────────────────────────────
TEMPLATE_FILE = 'Test.docx'

# ── Тексты характеристик ────────────────────────────────────────────────────────
CHAR_TEXTS = {
    'положительная': [
        'Постоянно проживает по адресу регистрации. По месту жительства характеризуется положительно, жалоб от соседей и родственников не поступало. На профилактические беседы является своевременно.',
        'По месту жительства характеризуется положительно. К законным требованиям сотрудников полиции относится корректно, установленный порядок соблюдает.',
        'В быту ведет себя спокойно, конфликтных ситуаций не создает. По месту жительства характеризуется удовлетворительно, замечаний не поступало.',
        'Поддерживает социально-бытовые связи с родственниками, в общественных местах ведет себя сдержанно. По месту жительства характеризуется положительно.',
        'За отчетный период нарушений общественного порядка не установлено. По месту жительства характеризуется положительно, на замечания реагирует адекватно.',
        'К профилактической работе относится добросовестно, на вызовы и беседы является без нарушений сроков. По месту жительства характеризуется удовлетворительно.',
        'По информации соседей и родственников характеризуется положительно. Склонности к конфликтному поведению не отмечается.',
        'В быту опрятен(а), общение с окружающими корректное. По месту жительства характеризуется удовлетворительно, отрицательных сведений не получено.',
        'Установленные обязанности выполняет, в ходе проверок замечаний не выявлено. По месту жительства характеризуется положительно.',
        'За период наблюдения существенных нарушений не допускал(а). По месту жительства характеризуется удовлетворительно, жалоб не поступало.',
    ],
    'нейтральная': [
        'Постоянно проживает по адресу регистрации. По месту жительства характеризуется удовлетворительно, выраженных положительных либо отрицательных сведений не поступало.',
        'По месту жительства характеризуется удовлетворительно. В общении с окружающими ведет себя нейтрально, грубых замечаний не имеется.',
        'На профилактические мероприятия является, к проводимой работе относится формально. По месту жительства характеризуется удовлетворительно.',
        'За отчетный период существенных нарушений общественного порядка не установлено. По месту жительства характеризуется нейтрально.',
        'В быту ведет себя без выраженных отклонений, конфликтных ситуаций не инициирует. По месту жительства характеризуется удовлетворительно.',
        'К установленным обязанностям относится без выраженной инициативы, но требования в целом соблюдает. По месту жительства характеризуется удовлетворительно.',
        'По информации соседей и родственников характеризуется удовлетворительно, отрицательных сведений, требующих немедленного реагирования, не поступало.',
        'Поведение в быту стабильное, нуждается в дальнейшем профилактическом контроле. По месту жительства характеризуется удовлетворительно.',
        'В общественных местах ведет себя сдержанно, на замечания реагирует. По месту жительства характеризуется нейтрально.',
        'По итогам проверок оценка поведения удовлетворительная: сохраняется необходимость планового профилактического наблюдения.',
    ],
    'отрицательная': [
        'По месту жительства характеризуется отрицательно. Имеются замечания к соблюдению установленного порядка, профилактические рекомендации выполняет не в полном объеме.',
        'На профилактические беседы является несвоевременно, к законным требованиям относится формально. По месту жительства характеризуется отрицательно.',
        'По информации соседей и родственников отмечаются конфликтные проявления в быту. По месту жительства характеризуется отрицательно.',
        'За отчетный период допускал(а) нарушения установленных обязанностей, выводы из замечаний делает не всегда. По месту жительства характеризуется отрицательно.',
        'К профилактической работе относится без должной ответственности, контроль требует усиления. По месту жительства характеризуется отрицательно.',
        'В быту поведение нестабильное, на замечания реагирует выборочно. По месту жительства характеризуется отрицательно.',
        'Допускает несвоевременное исполнение обязанностей, возложенных приговором суда. По месту жительства характеризуется отрицательно.',
        'По результатам проверок положительной динамики поведения не отмечается. По месту жительства характеризуется отрицательно.',
        'Имеются повторные замечания по соблюдению ограничений и установленного порядка. По месту жительства характеризуется отрицательно.',
        'По совокупности полученных сведений характеризуется отрицательно, требуется дальнейшая адресная профилактическая работа.',
    ],
}

CONN_TEXTS = {
    'положительная': (
        'Имеются сведения об отсутствии связей с лицами, ранее судимыми.'
    ),
    'нейтральная': 'Сведений об отсутсвии связей с ранее судимыми не имеются.',
    'отрицательная': 'Имеются сведения о связях с лицами, ранее судимыми.',
}

FEATURES_TEXT = 'Особых примет не выявлено.'
SEASON_CLOTHES_TEXT = 'Одет по сезону.'
IC_CHECK_TEXT = 'См. справку ИБД-Р'

QUARTER_ACC   = {1: 'первый', 2: 'второй', 3: 'третий', 4: 'четвёртый'}
CHAR_OPTIONS  = ['положительная', 'нейтральная', 'отрицательная']

_P12_TEMPLATE = (
    'Мной, {position} ГУУП 15 отдела полиции УМВД России по Калининскому '
    'району г. СПб {rank_instr} {fio_instr}, в ходе проведения профилактического '
    'обхода территории административного участка проверен (а) по месту '
    'жительства гражданин (ка):'
)
_P30_SPACES = ' ' * 57

# ── Творительный падеж ─────────────────────────────────────────────────────────

_RANK_INSTR: dict = {
    'капитан': 'капитаном', 'майор': 'майором', 'лейтенант': 'лейтенантом',
    'подполковник': 'подполковником', 'полковник': 'полковником',
    'старший': 'старшим', 'полиции': 'полиции', 'ст.': 'ст.',
}


def _rank_to_instr(rank: str) -> str:
    return ' '.join(_RANK_INSTR.get(w.lower(), w) for w in rank.split())


def _surname_to_instr(s: str) -> str:
    if re.search(r'(?:енко|(?<![ео])ко|[лн]о)$', s, re.I):
        return s
    if s.endswith('ский'): return s[:-4] + 'ским'
    if s.endswith('цкий'): return s[:-4] + 'цким'
    if s.endswith('ый'):   return s[:-2] + 'ым'
    if s.endswith('ий'):   return s[:-2] + 'им'
    if s.endswith('ой'):   return s[:-2] + 'ым'
    for end in ('ов', 'ев', 'ёв', 'ин', 'ын', 'ун', 'ан', 'он', 'ен'):
        if s.endswith(end): return s + 'ым'
    if s[-1] in ('ч', 'щ'):        return s + 'ем'
    if s[-1] in ('ж', 'ш', 'ц'):   return s + 'ом'
    if s.endswith('а'):  return s[:-1] + 'ой'
    if s.endswith('я'):  return s[:-1] + 'ей'
    return s + 'ом'


def _fio_initials(fio: str) -> str:
    parts = fio.strip().split()
    return ''.join(p[0].upper() + '.' for p in parts[1:] if p)


def _fio_display(fio: str) -> str:
    parts = fio.strip().split()
    if not parts: return fio
    initials = ''.join(p[0].upper() + '.' for p in parts[1:] if p)
    return f'{parts[0]} {initials}' if initials else parts[0]


def _fio_instr(fio: str) -> str:
    parts = fio.strip().split()
    if not parts: return fio
    return f'{_surname_to_instr(parts[0])} {_fio_initials(fio)}'.strip()


def _officer_label(officer: dict) -> str:
    """Метка вида 'Асташенков (40 / Вак)' или 'Кузьмин (17)'."""
    district = officer.get('district', '')
    fio      = officer.get('fio', '').strip()
    surname  = fio.split()[0] if fio else ''
    base = surname if surname else '—'
    if district and officer.get('is_vacancy'):
        return f'{base} ({district} / Вак)'
    if district:
        return f'{base} ({district})'
    if officer.get('is_vacancy'):
        return f'{base} (Вак)'
    return base


def _officer_folder(officer: dict) -> str:
    """Имя папки (безопасное для Windows) из метки участкового."""
    return re.sub(r'[/\\:*?"<>|]', '-', _officer_label(officer)).strip()


# ── Сопоставление адреса ───────────────────────────────────────────────────────

_STOP_ADDR = re.compile(
    r'\b(?:г|санкт.петербург|санкт|петербург|спб|рф|зарег|прож|'
    r'д|дом|кв|квартира|корп|корпус|пр.т|пр|ул|бул|б.р|наб|пл|ш|пер|'
    r'шоссе|проспект|улица|бульвар|набережная|площадь|переулок|'
    r'литер|лит|стр|строение|зд|здание)\b[.\s-]*',
    re.IGNORECASE,
)


def _norm(s: str) -> str:
    """Нижний регистр + ё→е (типичная разница в написании адресов)."""
    return s.replace('ё', 'е').replace('Ё', 'Е').lower()


def _extract_house(line: str):
    """Номер дома из строки адреса или None."""
    m = re.search(r'\bдд?\.?\s*(\d+)', line, re.IGNORECASE)
    if not m:
        m = re.search(r',\s*(\d+)\b', line)
    return m.group(1) if m else None


def _house_in_segment(house_str: str, segment: str) -> bool:
    """Проверяет, входит ли номер дома в сегмент адреса участкового.
    Обрабатывает точные номера, диапазоны «1–100» и списки «1, 3, 5»."""
    try:
        house = int(house_str)
    except ValueError:
        return house_str in segment

    # Диапазоны: «1-100», «1–100», «1—100»
    for m in re.finditer(r'(\d+)\s*[-–—]\s*(\d+)', segment):
        lo, hi = int(m.group(1)), int(m.group(2))
        if lo <= house <= hi:
            return True

    # Точное совпадение числа
    if re.search(r'\b' + re.escape(house_str) + r'\b', segment):
        return True

    return False


def _keywords_from(line: str) -> list:
    """Список значимых Cyrillic-слов (4+ букв) после удаления стоп-слов."""
    clean = _STOP_ADDR.sub(' ', line)
    skip  = {'квартира', 'корпус', 'этаж', 'литер', 'строение', 'здание'}
    return [w for w in re.findall(r'[А-ЯЁа-яё]{4,}', clean)
            if _norm(w) not in skip]


def _match_officer(address: str, officers: list) -> dict:
    if not address or not officers:
        return None

    lines = [l.strip() for l in address.split('\n') if len(l.strip()) > 8]
    if not lines:
        return None

    # Подготавливаем данные по каждой строке адреса
    line_data = []
    for line in lines[:3]:
        kws = _keywords_from(line)
        if kws:
            house = _extract_house(line)
            line_data.append((kws, house))

    if not line_data:
        return None

    # Проход 1: ключевое слово + номер дома (с поддержкой диапазонов)
    for kws, house in line_data:
        for kw in kws[:3]:
            kw_n = _norm(kw)
            for officer in officers:
                addrs_n = _norm(officer.get('addresses', ''))
                idx = addrs_n.find(kw_n)
                while idx >= 0:
                    seg = officer['addresses'][idx: idx + 400]
                    if house is None or _house_in_segment(house, seg):
                        return officer
                    idx = addrs_n.find(kw_n, idx + 1)

    # Проход 2: только улица (fallback) — однозначное совпадение без номера дома
    for kws, house in line_data:
        if house is None:
            continue   # без дома — пропускаем fallback (слишком неточно)
        candidates = []
        for kw in kws[:2]:
            kw_n = _norm(kw)
            for officer in officers:
                if kw_n in _norm(officer.get('addresses', '')) and officer not in candidates:
                    candidates.append(officer)
        if len(candidates) == 1:
            return candidates[0]

    return None


# ── docx-утилиты ───────────────────────────────────────────────────────────────

def _unique_cells(row) -> list:
    seen, result = set(), []
    for cell in row.cells:
        cid = id(cell._tc)
        if cid not in seen:
            seen.add(cid)
            result.append(cell)
    return result


def _set_cell_text(cell, text: str):
    if not cell.paragraphs:
        return
    para = cell.paragraphs[0]
    rPr_copy = None
    for r_elem in para._element.findall(qn('w:r')):
        rpr = r_elem.find(qn('w:rPr'))
        if rpr is not None:
            rPr_copy = copy.deepcopy(rpr)
        break
    for r_elem in list(para._element.findall(qn('w:r'))):
        para._element.remove(r_elem)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    r = OxmlElement('w:r')
    if rPr_copy is not None:
        r.append(rPr_copy)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    para._element.append(r)


def _set_para_text(para, text: str):
    rPr_copy = None
    for r_elem in para._element.findall(qn('w:r')):
        rpr = r_elem.find(qn('w:rPr'))
        if rpr is not None:
            rPr_copy = copy.deepcopy(rpr)
        break
    for r_elem in list(para._element.findall(qn('w:r'))):
        para._element.remove(r_elem)
    r = OxmlElement('w:r')
    if rPr_copy is not None:
        r.append(rPr_copy)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    para._element.append(r)


# ── Генерация одной справки ────────────────────────────────────────────────────

def _pick_char_template(char_type: str, char_templates: dict = None) -> str:
    source = char_templates if isinstance(char_templates, dict) else CHAR_TEXTS
    raw = source.get(char_type)
    if raw is None:
        raw = source.get('нейтральная', '')
    if isinstance(raw, str):
        return raw
    if isinstance(raw, (list, tuple)):
        options = [x.strip() for x in raw if isinstance(x, str) and x.strip()]
        if options:
            return random.choice(options)
    return ''


def generate_one(rec: dict, char_type: str, quarter: int, year: int,
                 out_dir: str, officer: dict = None, custom_char_text: str = "",
                 char_templates: dict = None) -> str:
    doc    = Document(TEMPLATE_FILE)
    tables = doc.tables

    fio      = rec.get('ФИО', '').strip()
    dr       = rec.get('Дата рождения', '').strip()
    court    = rec.get('Суд (когда, кем)', '').strip()
    duties   = rec.get('Обязанности', '').strip()
    end_date = rec.get('Окончание срока', '').strip()
    address  = rec.get('Место жительства', '').strip()
    work     = rec.get('Место работы (учебы)', '').strip()
    phone    = rec.get('Телефон', '').strip()
    manual_char_text = (rec.get('Характеристика') or rec.get('Характеристика (п.8)', '')).strip()
    manual_links_text = (rec.get('Связи') or rec.get('Связи (п.9)', '')).strip()
    manual_features_text = (rec.get('Приметы') or rec.get('Приметы (п.10)', '')).strip()
    manual_season_text = (rec.get('Сезонная одежда') or rec.get('Сезонная одежда (п.11)', '')).strip()
    manual_violations_text = (rec.get('Нарушения') or rec.get('Нарушения (п.12)', '')).strip()
    manual_ic_check_text = (rec.get('Проверка ИЦ') or rec.get('Проверка ИЦ (п.13)', '')).strip()

    # User-editable template selected randomly per document.
    custom_char_text = (custom_char_text or '').strip()
    if custom_char_text:
        char_text = custom_char_text
        conn_text = CONN_TEXTS.get('нейтральная', '')
    else:
        safe_char = char_type if char_type in CHAR_OPTIONS else 'нейтральная'
        char_text = _pick_char_template(safe_char, char_templates=char_templates)
        conn_text = CONN_TEXTS.get(safe_char, CONN_TEXTS.get('нейтральная', ''))

    if manual_char_text:
        char_text = manual_char_text
    if manual_links_text:
        conn_text = manual_links_text
    features_text = manual_features_text or FEATURES_TEXT
    season_text = manual_season_text or SEASON_CLOTHES_TEXT
    violations_text = manual_violations_text
    ic_check_text = manual_ic_check_text or IC_CHECK_TEXT

    def safe_t(idx):
        return tables[idx] if idx < len(tables) else None

    t0 = safe_t(0)
    if t0 and len(t0.rows) > 1:
        uc = _unique_cells(t0.rows[0])
        if len(uc) > 1: _set_cell_text(uc[1], fio)
        uc = _unique_cells(t0.rows[1])
        if len(uc) > 1: _set_cell_text(uc[1], dr)

    t1 = safe_t(1)
    if t1 and t1.rows:
        uc = _unique_cells(t1.rows[0])
        if uc:
            cell, old = uc[0], uc[0].text
            prefix = (old[: old.index(': ') + 2] if ': ' in old
                      else old[: old.index(':') + 2] if ':' in old
                      else '3. Осуждён(а): ')
            _set_cell_text(cell, prefix + court)

    for t_idx, col in ((2, 1), (3, 1), (4, 1)):
        t = safe_t(t_idx)
        if t and t.rows:
            uc = _unique_cells(t.rows[0])
            if len(uc) > col:
                _set_cell_text(uc[col], [duties, end_date, address][t_idx - 2])

    t5 = safe_t(5)
    if t5 and t5.rows:
        uc = _unique_cells(t5.rows[0])
        if len(uc) > 1:
            _set_cell_text(uc[1], work)
        if len(uc) > 3:
            _set_cell_text(uc[3], phone)

    t6 = safe_t(6)
    if t6 and len(t6.rows) > 2:
        uc = _unique_cells(t6.rows[2])
        if uc: _set_cell_text(uc[0], char_text)

    t7 = safe_t(7)
    if t7 and t7.rows:
        uc = _unique_cells(t7.rows[0])
        if len(uc) > 1: _set_cell_text(uc[1], conn_text)
        if len(t7.rows) > 1:
            uc = _unique_cells(t7.rows[1])
            if len(uc) > 1:
                _set_cell_text(uc[1], features_text)

    t8 = safe_t(8)
    if t8 and t8.rows:
        uc = _unique_cells(t8.rows[0])
        if len(uc) > 1:
            _set_cell_text(uc[1], season_text)

    t9 = safe_t(9)
    if t9 and t9.rows:
        uc = _unique_cells(t9.rows[0])
        if len(uc) > 1:
            _set_cell_text(uc[1], violations_text)

    t10 = safe_t(10)
    if t10 and t10.rows:
        uc = _unique_cells(t10.rows[0])
        if len(uc) > 1:
            _set_cell_text(uc[1], ic_check_text)

    new_period = f'(за {quarter} квартал {year} года)'
    for para in doc.paragraphs:
        if 'квартал' in para.text.lower():
            _set_para_text(para, new_period)
            break

    if officer:
        pos   = officer.get('position', 'УУП')
        rank  = officer.get('rank', '')
        fio_o = officer.get('fio', '')
        parts = fio_o.strip().split()
        inits = ''.join(p[0].upper() + '.' for p in parts[1:] if p)
        sur   = parts[0] if parts else ''
        paras = doc.paragraphs
        if len(paras) > 12:
            _set_para_text(paras[12], _P12_TEMPLATE.format(
                position=pos, rank_instr=_rank_to_instr(rank), fio_instr=_fio_instr(fio_o)))
        if len(paras) > 28:
            _set_para_text(paras[28], f'{pos} ГУУП 15 отдела полиции')
        if len(paras) > 30:
            _set_para_text(paras[30], f'{rank}{_P30_SPACES}{inits} {sur}')

    os.makedirs(out_dir, exist_ok=True)
    safe_name = re.sub(r'[\\/:*?"<>|]', '_', fio).strip() or 'без_имени'
    out_path  = os.path.abspath(os.path.join(out_dir, f'{safe_name}.docx'))
    doc.save(out_path)
    return out_path
